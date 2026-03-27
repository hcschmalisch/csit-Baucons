from flask import Flask, request, jsonify, redirect
import base64
import tempfile
import os
import logging
import json
from time import sleep
from datetime import datetime
from io import BytesIO
from azure.identity import DefaultAzureCredential
from azure.core.credentials import AzureKeyCredential
from azure.ai.projects import AIProjectClient
from functools import wraps
import threading
import multiprocessing
import uuid
from typing import Dict, Any
import filelock

# Logging konfigurieren
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('api_server.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================
# Azure AD Authentication Decorator
# ============================================

def require_auth(f):
    """
    Decorator für Azure AD Authentication (Easy Auth)
    Schützt nur Admin-Routen, API-Endpoints bleiben mit API-Keys erreichbar
    """
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Easy Auth setzt diesen Header wenn User angemeldet ist
        user_principal = request.headers.get('X-MS-CLIENT-PRINCIPAL')
        if not user_principal:
            # Redirect zur Azure AD Login-Seite
            return redirect('/.auth/login/aad?post_login_redirect_uri=' + request.path)
        return f(*args, **kwargs)
    return decorated_function

# ============================================
# Usage Tracking für Abrechnung
# ============================================

USAGE_LOG_FILE = os.path.join(tempfile.gettempdir(), 'api_usage_log.json')
USAGE_LOCK_FILE = USAGE_LOG_FILE + '.lock'

def log_usage(customer_name: str, request_id: str, file_names: list, file_sizes: list, 
              success: bool, tokens_input: int = 0, tokens_output: int = 0, 
              action: str = None, error_message: str = None):
    """
    Loggt die API-Nutzung für Abrechnungszwecke
    
    Parameter:
    - customer_name: Name des Kunden
    - request_id: Eindeutige Request-ID
    - file_names: Liste der Dateinamen
    - file_sizes: Liste der Dateigrößen in Bytes
    - success: Ob die Anfrage erfolgreich war
    - tokens_input: Anzahl der Input-Tokens
    - tokens_output: Anzahl der Output-Tokens
    - action: Die ausgeführte Aktion
    - error_message: Fehlermeldung bei Misserfolg
    """
    now = datetime.now()
    
    usage_entry = {
        "timestamp": now.isoformat(),
        "date": now.strftime("%Y-%m-%d"),
        "time": now.strftime("%H:%M:%S"),
        "request_id": request_id,
        "customer": customer_name,
        "action": action,
        "success": success,
        "files": [
            {"name": name, "size_bytes": size} 
            for name, size in zip(file_names, file_sizes)
        ],
        "total_file_size_bytes": sum(file_sizes),
        "tokens": {
            "input": tokens_input,
            "output": tokens_output,
            "total": tokens_input + tokens_output
        },
        "error": error_message
    }
    
    # Thread-safe Schreiben mit File Lock
    try:
        lock = filelock.FileLock(USAGE_LOCK_FILE, timeout=10)
        with lock:
            # Existierende Daten laden
            usage_data = {"entries": []}
            if os.path.exists(USAGE_LOG_FILE):
                try:
                    with open(USAGE_LOG_FILE, 'r', encoding='utf-8') as f:
                        usage_data = json.load(f)
                except (json.JSONDecodeError, IOError):
                    usage_data = {"entries": []}
            
            # Neuen Eintrag hinzufügen
            usage_data["entries"].append(usage_entry)
            
            # Speichern
            with open(USAGE_LOG_FILE, 'w', encoding='utf-8') as f:
                json.dump(usage_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"[{request_id}] Usage logged: {customer_name}, {len(file_names)} files, "
                   f"{tokens_input + tokens_output} tokens, success={success}")
    except Exception as e:
        logger.error(f"[{request_id}] Failed to log usage: {e}")

def get_usage_summary(customer_name: str = None, start_date: str = None, end_date: str = None) -> dict:
    """
    Gibt eine Zusammenfassung der Nutzung zurück
    
    Parameter:
    - customer_name: Optional - Filter nach Kunde
    - start_date: Optional - Startdatum (YYYY-MM-DD)
    - end_date: Optional - Enddatum (YYYY-MM-DD)
    """
    try:
        if not os.path.exists(USAGE_LOG_FILE):
            return {"entries": [], "summary": {}}
        
        with open(USAGE_LOG_FILE, 'r', encoding='utf-8') as f:
            usage_data = json.load(f)
        
        entries = usage_data.get("entries", [])
        
        # Filter anwenden
        if customer_name:
            entries = [e for e in entries if e.get("customer") == customer_name]
        if start_date:
            entries = [e for e in entries if e.get("date", "") >= start_date]
        if end_date:
            entries = [e for e in entries if e.get("date", "") <= end_date]
        
        # Zusammenfassung berechnen
        summary = {}
        for entry in entries:
            customer = entry.get("customer", "unknown")
            if customer not in summary:
                summary[customer] = {
                    "total_requests": 0,
                    "successful_requests": 0,
                    "failed_requests": 0,
                    "total_files": 0,
                    "total_file_size_bytes": 0,
                    "total_tokens_input": 0,
                    "total_tokens_output": 0,
                    "total_tokens": 0,
                    "by_action": {}
                }
            
            s = summary[customer]
            s["total_requests"] += 1
            if entry.get("success"):
                s["successful_requests"] += 1
            else:
                s["failed_requests"] += 1
            s["total_files"] += len(entry.get("files", []))
            s["total_file_size_bytes"] += entry.get("total_file_size_bytes", 0)
            tokens = entry.get("tokens", {})
            s["total_tokens_input"] += tokens.get("input", 0)
            s["total_tokens_output"] += tokens.get("output", 0)
            s["total_tokens"] += tokens.get("total", 0)
            
            # Nach Aktion gruppieren
            action = entry.get("action", "unknown")
            if action not in s["by_action"]:
                s["by_action"][action] = {"requests": 0, "tokens": 0}
            s["by_action"][action]["requests"] += 1
            s["by_action"][action]["tokens"] += tokens.get("total", 0)
        
        return {"entries": entries, "summary": summary}
    except Exception as e:
        logger.error(f"Failed to get usage summary: {e}")
        return {"error": str(e)}

# ============================================
# Konvertierungsfunktionen
# ============================================

def convert_to_pdf(file_content, file_extension, original_filename, request_id):
    """
    Konvertiert DOCX oder XLSX Dateien zu PDF
    """
    if file_extension == '.docx':
        return convert_docx_to_pdf(file_content, request_id)
    elif file_extension == '.xlsx':
        return convert_xlsx_to_pdf(file_content, original_filename, request_id)
    else:
        raise ValueError(f"Nicht unterstützter Dateityp: {file_extension}")

def convert_docx_to_pdf(docx_content, request_id):
    """
    Konvertiert DOCX zu PDF mit python-docx und reportlab
    """
    try:
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        import html
        
        # DOCX laden
        doc = Document(BytesIO(docx_content))
        
        # PDF erstellen
        pdf_buffer = BytesIO()
        pdf_doc = SimpleDocTemplate(
            pdf_buffer, 
            pagesize=A4,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        story = []
        styles = getSampleStyleSheet()
        
        # Erstelle einen benutzerdefinierten Stil für normalen Text
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            alignment=TA_LEFT
        )
        
        # Zähle Paragraphen und Tabellen für Logging
        para_count = 0
        table_count = 0
        
        # Durchlaufe alle Paragraphen
        for para in doc.paragraphs:
            if para.text.strip():
                para_count += 1
                # Escape HTML-Sonderzeichen und entferne problematische Zeichen
                text = html.escape(para.text)
                # Entferne NULL-Bytes und andere Steuerzeichen
                text = ''.join(char for char in text if char.isprintable() or char in ['\n', '\r', '\t'])
                
                try:
                    # Bestimme Stil basierend auf Formatierung
                    if para.style.name.startswith('Heading'):
                        story.append(Paragraph(text, styles['Heading2']))
                    else:
                        story.append(Paragraph(text, normal_style))
                    story.append(Spacer(1, 0.15*inch))
                except Exception as para_error:
                    # Fallback: Füge als Plain-Text hinzu wenn Paragraph fehlschlägt
                    logger.warning(f"[{request_id}] Paragraph-Fehler, verwende Plain-Text: {str(para_error)[:100]}")
                    story.append(Spacer(1, 0.15*inch))
        
        # Durchlaufe alle Tabellen
        for table in doc.tables:
            table_count += 1
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Escape und bereinige Zelltext
                    cell_text = html.escape(cell.text)
                    cell_text = ''.join(char for char in cell_text if char.isprintable() or char in ['\n', '\r', '\t'])
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data:
                try:
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.25*inch))
                except Exception as table_error:
                    logger.warning(f"[{request_id}] Tabellen-Fehler: {str(table_error)[:100]}")
        
        # Log Statistiken
        logger.info(f"[{request_id}] DOCX Inhalt: {para_count} Paragraphen, {table_count} Tabellen")
        
        if not story:
            logger.warning(f"[{request_id}] DOCX enthält keinen konvertierbaren Inhalt!")
            # Erstelle eine Fehlermeldung im PDF
            story.append(Paragraph("WARNUNG: Dokument enthält keinen lesbaren Text.", styles['Normal']))
        
        pdf_doc.build(story)
        logger.info(f"[{request_id}] DOCX erfolgreich zu PDF konvertiert ({len(pdf_buffer.getvalue())} Bytes)")
        return pdf_buffer.getvalue()
    except Exception as e:
        logger.error(f"[{request_id}] DOCX Konvertierung fehlgeschlagen: {e}")
        raise

def convert_xlsx_to_pdf(xlsx_content, original_filename, request_id):
    """
    Konvertiert XLSX zu PDF mit openpyxl und reportlab.
    Unterstützt Merged Cells, Spaltenbreiten, dynamische Font-Größe und Textumbruch.
    """
    try:
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, PageBreak, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, mm

        wb = load_workbook(BytesIO(xlsx_content), data_only=True)

        pdf_buffer = BytesIO()
        page_w, page_h = landscape(A4)
        margin = 15 * mm
        usable_w = page_w - 2 * margin
        pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=landscape(A4),
                                    leftMargin=margin, rightMargin=margin,
                                    topMargin=margin, bottomMargin=margin)
        story = []
        styles = getSampleStyleSheet()

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # --- Merged Cells erfassen (nur Wert in die erste Zelle, Rest leer) ---
            merged_first_cells = {}  # (row, col) -> value (nur linke obere Zelle)
            merged_skip_cells = set()  # Alle anderen Zellen im Merge-Bereich
            for merged_range in list(ws.merged_cells.ranges):
                mc_min_col, mc_min_row, mc_max_col, mc_max_row = merged_range.bounds
                top_left_value = ws.cell(row=mc_min_row, column=mc_min_col).value
                merged_first_cells[(mc_min_row, mc_min_col)] = top_left_value
                for r in range(mc_min_row, mc_max_row + 1):
                    for c in range(mc_min_col, mc_max_col + 1):
                        if (r, c) != (mc_min_row, mc_min_col):
                            merged_skip_cells.add((r, c))

            # --- Zellenwerte lesen ---
            min_row = ws.min_row or 1
            max_row = ws.max_row or 1
            min_col = ws.min_column or 1
            max_col = ws.max_column or 1

            raw_data = []
            for r in range(min_row, max_row + 1):
                row_vals = []
                for c in range(min_col, max_col + 1):
                    if (r, c) in merged_skip_cells:
                        row_vals.append('')
                    elif (r, c) in merged_first_cells:
                        val = merged_first_cells[(r, c)]
                        row_vals.append(str(val) if val is not None else '')
                    else:
                        val = ws.cell(row=r, column=c).value
                        row_vals.append(str(val) if val is not None else '')
                raw_data.append(row_vals)

            if not raw_data:
                continue

            # --- Spalten mit tatsächlichen Daten identifizieren ---
            # Zähle pro Spalte wie viele Zeilen Daten haben (Merged-Header zählen nur 1x)
            num_cols = len(raw_data[0]) if raw_data else 0
            col_data_count = [0] * num_cols
            for row in raw_data:
                for ci, val in enumerate(row):
                    if val.strip():
                        col_data_count[ci] += 1

            # Behalte nur Spalten die in mindestens 2 Zeilen Daten haben
            # (filtert einzelne Merged-Header-Zellen die über viele Spalten gehen)
            keep_cols = [i for i, cnt in enumerate(col_data_count) if cnt >= 2]
            if not keep_cols:
                # Fallback: alle Spalten mit min. 1 Datenpunkt
                keep_cols = [i for i, cnt in enumerate(col_data_count) if cnt >= 1]
            if not keep_cols:
                continue

            filtered_data = []
            for row in raw_data:
                new_row = [row[i] for i in keep_cols]
                if any(v.strip() for v in new_row):
                    filtered_data.append(new_row)

            if not filtered_data:
                continue

            # --- Spaltenbreiten aus XLSX übernehmen ---
            col_widths_raw = []
            for orig_ci in keep_cols:
                col_letter = get_column_letter(min_col + orig_ci)
                dim = ws.column_dimensions.get(col_letter)
                w = dim.width if dim and dim.width else 10
                col_widths_raw.append(max(w, 3))

            total_raw = sum(col_widths_raw)
            col_widths = [(w / total_raw) * usable_w for w in col_widths_raw]

            # --- Dynamische Font-Größe basierend auf Spaltenanzahl ---
            num_visible_cols = len(keep_cols)
            if num_visible_cols <= 4:
                font_size = 9
            elif num_visible_cols <= 7:
                font_size = 8
            else:
                font_size = 7

            cell_style = ParagraphStyle('CellStyle', parent=styles['Normal'],
                                        fontSize=font_size, leading=font_size + 2)
            header_style = ParagraphStyle('HeaderStyle', parent=styles['Normal'],
                                          fontSize=font_size, leading=font_size + 2,
                                          textColor=colors.white)

            # Sheet-Titel
            story.append(Paragraph(f"<b>{sheet_name}</b>", styles['Heading2']))
            story.append(Spacer(1, 3 * mm))

            # --- Tabellendaten mit Paragraph (Textumbruch) ---
            table_data = []
            for row in filtered_data[:150]:
                para_row = [Paragraph(v.replace('&', '&amp;').replace('<', '&lt;'), cell_style)
                            for v in row]
                table_data.append(para_row)

            if not table_data:
                continue

            t = Table(table_data, colWidths=col_widths, repeatRows=1)

            # --- Header-Zeile erkennen (erste Zeile mit vielen gefüllten Zellen) ---
            header_row_idx = 0
            for ri, row in enumerate(filtered_data[:10]):
                filled = sum(1 for v in row if v.strip())
                if filled >= num_visible_cols * 0.6:
                    header_row_idx = ri
                    # Header-Zellen mit header_style neu erstellen
                    table_data[ri] = [Paragraph(v.replace('&', '&amp;').replace('<', '&lt;'), header_style)
                                      for v in filtered_data[ri]]
                    break

            # --- Styling ---
            style_cmds = [
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTSIZE', (0, 0), (-1, -1), font_size),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.Color(0.6, 0.6, 0.6)),
                # Header-Zeile
                ('BACKGROUND', (0, header_row_idx), (-1, header_row_idx), colors.Color(0.3, 0.3, 0.3)),
                ('TEXTCOLOR', (0, header_row_idx), (-1, header_row_idx), colors.whitesmoke),
            ]

            # Alternierende Zeilenfarben (nach Header)
            for ri in range(header_row_idx + 1, len(table_data)):
                if (ri - header_row_idx) % 2 == 0:
                    style_cmds.append(('BACKGROUND', (0, ri), (-1, ri), colors.Color(0.95, 0.95, 0.95)))

            t.setStyle(TableStyle(style_cmds))
            story.append(t)

            if len(filtered_data) > 150:
                story.append(Spacer(1, 3 * mm))
                story.append(Paragraph(f"<i>... ({len(filtered_data) - 150} weitere Zeilen)</i>", styles['Normal']))

            story.append(PageBreak())

        pdf_doc.build(story)
        logger.info(f"[{request_id}] XLSX erfolgreich zu PDF konvertiert ({len(wb.sheetnames)} Sheets)")
        return pdf_buffer.getvalue()
    except Exception as e:
        logger.error(f"[{request_id}] XLSX Konvertierung fehlgeschlagen: {e}")
        raise

# ============================================
# Asynchrone Job-Verwaltung
# ============================================

# File-based Storage für Jobs (persistent über Worker-Instanzen)
JOBS_DIR = os.path.join(tempfile.gettempdir(), 'api_jobs')
os.makedirs(JOBS_DIR, exist_ok=True)
jobs_lock = threading.Lock()

def get_job_file_path(task_id: str) -> str:
    """Gibt den Pfad zur Job-Datei zurück"""
    return os.path.join(JOBS_DIR, f"{task_id}.json")

def create_job(task_id: str):
    """Erstellt einen neuen Job"""
    job_data = {
        'task_id': task_id,
        'status': 'pending',  # pending, processing, completed, failed
        'created_at': datetime.now().isoformat(),
        'started_at': None,
        'completed_at': None,
        'result': None,
        'error': None,
        'progress': 0
    }
    with jobs_lock:
        with open(get_job_file_path(task_id), 'w') as f:
            json.dump(job_data, f)

def update_job(task_id: str, **kwargs):
    """Aktualisiert einen Job"""
    with jobs_lock:
        job_path = get_job_file_path(task_id)
        if os.path.exists(job_path):
            with open(job_path, 'r') as f:
                job_data = json.load(f)
            job_data.update(kwargs)
            with open(job_path, 'w') as f:
                json.dump(job_data, f)

def get_job(task_id: str) -> Dict[str, Any]:
    """Holt Job-Informationen"""
    job_path = get_job_file_path(task_id)
    if os.path.exists(job_path):
        with open(job_path, 'r') as f:
            return json.load(f)
    return None

def cleanup_old_jobs():
    """Entfernt Jobs älter als 24 Stunden"""
    try:
        now = datetime.now()
        for filename in os.listdir(JOBS_DIR):
            if filename.endswith('.json'):
                filepath = os.path.join(JOBS_DIR, filename)
                try:
                    with open(filepath, 'r') as f:
                        job_data = json.load(f)
                    created = datetime.fromisoformat(job_data['created_at'])
                    if (now - created).total_seconds() > 86400:  # 24 Stunden
                        os.remove(filepath)
                except:
                    pass
    except:
        pass

def process_async_job(task_id: str, data: dict):
    """Verarbeitet einen Job im Hintergrund (läuft in eigenem Prozess)"""
    try:
        # Cache leeren: Child-Prozess erbt Sockets vom Parent die hier ungültig sind
        global _credential, _client_cache, _agent_cache
        _credential = None
        _client_cache = {}
        _agent_cache = {}

        update_job(task_id, status='processing', started_at=datetime.now().isoformat(), progress=10)

        # Führe die Analyse durch (gleiche Logik wie synchroner Endpoint)
        result = analyze_document_internal(data, task_id)
        
        update_job(
            task_id, 
            status='completed', 
            completed_at=datetime.now().isoformat(),
            result=result,
            progress=100
        )
        logger.info(f"[{task_id}] Async-Job erfolgreich abgeschlossen")
        
    except Exception as e:
        logger.error(f"[{task_id}] Async-Job fehlgeschlagen: {e}")
        update_job(
            task_id,
            status='failed',
            completed_at=datetime.now().isoformat(),
            error=str(e),
            progress=100
        )

# ============================================
# Flask App
# ============================================

app = Flask(__name__)

# Azure AI Configuration - Standardwerte
DEFAULT_ENDPOINT = "https://csit-brancons-resource.services.ai.azure.com/api/projects/csit-brancons"
DEFAULT_AGENT = "analyze-document"

# Azure AI Authentifizierung: Nur Managed Identity wird unterstützt
# Azure AI Projects SDK unterstützt nur TokenCredential (nicht AzureKeyCredential)
logger.info("✓ Azure AI Authentifizierung: Managed Identity (DefaultAzureCredential)")

# Globaler Credential & Client-Cache (einmalig erstellt, wiederverwendet pro Request)
_credential = None
_client_cache = {}  # key: endpoint -> {"project_client": ..., "openai_client": ...}
_agent_cache = {}   # key: (endpoint, agent_name) -> agent object
_cache_lock = threading.Lock()

def get_credential():
    """Gibt die gecachte DefaultAzureCredential zurück (einmalig erstellt)"""
    global _credential
    if _credential is None:
        logger.info("Erstelle DefaultAzureCredential (einmalig)...")
        _credential = DefaultAzureCredential()
        logger.info("✓ DefaultAzureCredential erstellt")
    return _credential

def get_clients(endpoint: str, force_refresh: bool = False):
    """Gibt gecachte AIProjectClient und OpenAI-Client für einen Endpoint zurück"""
    if force_refresh or endpoint not in _client_cache:
        with _cache_lock:
            if force_refresh or endpoint not in _client_cache:
                logger.info(f"Erstelle Clients für Endpoint: {endpoint}{' (refresh)' if force_refresh else ' (einmalig)'}...")
                credential = get_credential()
                import httpx
                project_client = AIProjectClient(endpoint=endpoint, credential=credential)
                openai_client = project_client.get_openai_client()
                # Kürzerer Connect-Timeout damit stale Connections schnell erkannt werden
                # Read-Timeout bleibt hoch (600s) für lange Agent-Aufrufe
                openai_client.timeout = httpx.Timeout(600.0, connect=10.0)
                openai_client.max_retries = 0  # Wir machen Retries selbst (mit Cache-Invalidierung)
                _client_cache[endpoint] = {
                    "project_client": project_client,
                    "openai_client": openai_client
                }
                logger.info(f"✓ Clients für {endpoint} erstellt und gecacht")
    return _client_cache[endpoint]["project_client"], _client_cache[endpoint]["openai_client"]

def invalidate_clients(endpoint: str):
    """Invalidiert den Client-Cache für einen Endpoint (bei Connection-Errors)"""
    with _cache_lock:
        _client_cache.pop(endpoint, None)
        keys_to_remove = [k for k in _agent_cache if k[0] == endpoint]
        for k in keys_to_remove:
            _agent_cache.pop(k, None)
    logger.info(f"Client-Cache invalidiert für {endpoint}")

def get_agent(endpoint: str, agent_name: str):
    """Gibt gecachten Agent zurück"""
    cache_key = (endpoint, agent_name)
    if cache_key not in _agent_cache:
        with _cache_lock:
            if cache_key not in _agent_cache:
                logger.info(f"Lade Agent '{agent_name}' (einmalig)...")
                project_client, _ = get_clients(endpoint)
                _agent_cache[cache_key] = project_client.agents.get(agent_name=agent_name)
                logger.info(f"✓ Agent '{agent_name}' geladen und gecacht")
    return _agent_cache[cache_key]

# Timeout-Konfiguration
# WICHTIG: Der Azure OpenAI Agent kann bei großen/komplexen Dateien mehrere Minuten benötigen
# Wenn Ihre Clients in einen Timeout laufen, erhöhen Sie den Timeout auf Client-Seite
# Der Server selbst hat keinen Timeout und wartet auf die Agent-Antwort
logger.info("⚠ HINWEIS: Azure Agent-Aufrufe können 1-5+ Minuten dauern")
logger.info("⚠ Client sollte min. 10 Minuten Timeout haben")

# Kundenkonfiguration laden
CUSTOMER_CONFIG = {}
CUSTOMER_CONFIG_FILE = os.path.join(os.path.dirname(__file__), 'customer_config.json')

def load_customer_config():
    """Lädt die Kundenkonfiguration aus der JSON-Datei"""
    global CUSTOMER_CONFIG
    try:
        logger.info(f"Lade Kundenkonfiguration von: {CUSTOMER_CONFIG_FILE}")
        with open(CUSTOMER_CONFIG_FILE, 'r', encoding='utf-8') as f:
            config_data = json.load(f)
            CUSTOMER_CONFIG = config_data.get('customers', {})
            if CUSTOMER_CONFIG is None:
                CUSTOMER_CONFIG = {}
            logger.info(f"✓ Kundenkonfiguration geladen: {len(CUSTOMER_CONFIG)} Kunden")
    except FileNotFoundError:
        logger.error(f"❌ customer_config.json nicht gefunden: {CUSTOMER_CONFIG_FILE}")
        logger.warning("⚠ Server läuft ohne Kundenauthentifizierung!")
    except Exception as e:
        logger.error(f"❌ Fehler beim Laden der Kundenkonfiguration: {e}")
        logger.warning("⚠ Server läuft ohne Kundenauthentifizierung!")

# Initial laden
load_customer_config()

def validate_customer_and_action(api_key, action):
    """
    Validiert API-Key und Aktion.
    Returns: (valid, customer_info, error_message)
    """
    if not api_key:
        return False, None, "API-Key fehlt"
    
    if api_key not in CUSTOMER_CONFIG:
        return False, None, "Ungültiger API-Key"
    
    customer = CUSTOMER_CONFIG[api_key]
    
    if not customer.get('active', False):
        return False, None, "Kunde ist deaktiviert"
    
    if action not in customer.get('allowed_actions', {}):
        return False, None, f"Aktion '{action}' nicht erlaubt für diesen Kunden"
    
    return True, customer, None

def get_endpoint_and_agent(api_key, action):
    """
    Ermittelt Endpoint und Agent basierend auf Kunde und Aktion.
    Returns: (endpoint, agent, customer_info) or (None, None, error_message)
    """
    valid, customer, error = validate_customer_and_action(api_key, action)
    
    if not valid:
        return None, None, error
    
    action_config = customer['allowed_actions'][action]
    endpoint = action_config['endpoint']
    agent = action_config['agent']
    
    return endpoint, agent, customer

logger.info("=" * 60)
logger.info("API Server wird gestartet...")
logger.info(f"Standard-Endpoint: {DEFAULT_ENDPOINT}")
logger.info(f"Standard-Agent: {DEFAULT_AGENT}")
logger.info(f"Authentifizierung: {'Aktiv' if CUSTOMER_CONFIG else 'Deaktiviert'}")
logger.info("=" * 60)

def analyze_document_internal(data: dict, request_id: str) -> dict:
    """
    Interne Funktion für Dokumentenanalyse (synchron und asynchron verwendbar)
    
    Unterstützt:
    - 0 Dateien (nur user_text)
    - 1 Datei (file_content/file_name ODER files-Array mit 1 Element)
    - N Dateien (files-Array)
    
    Wirft Exceptions bei Fehlern
    """
    uploaded_files = []
    temp_file_paths = []
    processed_file_names = []
    processed_file_sizes = []  # Für Usage Tracking
    customer_name = "unknown"
    action = None
    tokens_input = 0
    tokens_output = 0
    
    try:
        # Basis-Parameter
        api_key = data.get('api_key')
        action = data.get('action')
        user_text = data.get('user_text')
        vector_store_ids = data.get('vector_store_ids') or data.get('vector_store_id') or data.get('vectorstoreId') or data.get('vectorStoreId')
        
        # Vector Store IDs normalisieren
        if isinstance(vector_store_ids, str):
            vector_store_ids = [vector_store_ids]
        elif vector_store_ids is None:
            vector_store_ids = []
        elif not isinstance(vector_store_ids, list):
            raise ValueError("vector_store_ids muss eine Liste oder ein String sein")
        vector_store_ids = [vs_id for vs_id in vector_store_ids if isinstance(vs_id, str) and vs_id.strip()]

        # Dateien normalisieren: Alles in ein files-Array überführen
        files_array = data.get('files', [])
        
        # Rückwärtskompatibilität: file_content/file_name als einzelne Datei
        file_content_base64 = data.get('file_content')
        file_name = data.get('file_name')
        if file_content_base64 and file_name:
            files_array.append({
                'file_content': file_content_base64,
                'file_name': file_name
            })
        
        # Validierung: Mindestens Dateien ODER Text
        has_files = len(files_array) > 0
        has_text = bool(user_text)
        
        if not has_files and not has_text:
            raise ValueError("Mindestens eine Datei (files) oder user_text muss angegeben werden")
        
        # API-Key und Aktion Validierung
        if CUSTOMER_CONFIG:
            if not api_key:
                raise ValueError("api_key ist erforderlich")
            if not action:
                raise ValueError("action ist erforderlich")
            
            endpoint, agent_name, customer = get_endpoint_and_agent(api_key, action)
            if endpoint is None:
                raise ValueError(f"Autorisierung fehlgeschlagen: {agent_name}")
            
            customer_name = customer['customer_name']  # Für Usage Tracking
            logger.info(f"[{request_id}] Kunde: {customer_name}")
            logger.info(f"[{request_id}] Aktion: {action}")

            # Vector Store pro Kunde/Aktion (falls nicht im Request angegeben)
            if not vector_store_ids:
                action_config = customer.get('allowed_actions', {}).get(action, {})
                config_vector_store_ids = action_config.get('vector_store_ids') or action_config.get('vector_store_id')
                if isinstance(config_vector_store_ids, str):
                    config_vector_store_ids = [config_vector_store_ids]
                elif config_vector_store_ids is None:
                    config_vector_store_ids = []
                elif not isinstance(config_vector_store_ids, list):
                    raise ValueError("vector_store_ids in customer_config.json muss eine Liste oder ein String sein")
                vector_store_ids = [vs_id for vs_id in config_vector_store_ids if isinstance(vs_id, str) and vs_id.strip()]
            
            if 'endpoint' in data:
                endpoint = data['endpoint']
            if 'agent' in data:
                agent_name = data['agent']
        else:
            endpoint = data.get('endpoint', DEFAULT_ENDPOINT)
            agent_name = data.get('agent', DEFAULT_AGENT)
        
        logger.info(f"[{request_id}] Endpoint: {endpoint}")
        logger.info(f"[{request_id}] Agent: {agent_name}")
        logger.info(f"[{request_id}] Anzahl Dateien: {len(files_array)}")
        
        # Alle Dateien verarbeiten
        import re
        for idx, file_info in enumerate(files_array):
            f_content_raw = file_info.get('file_content', '')
            f_name = file_info.get('file_name', f'file_{idx}.bin')

            # Power Automate sendet file_content als Objekt: {"$content-type": "...", "$content": "base64..."}
            if isinstance(f_content_raw, dict):
                f_content_base64 = f_content_raw.get('$content', '')
                logger.info(f"[{request_id}] Power Automate Format erkannt (content-type: {f_content_raw.get('$content-type', 'unbekannt')})")
            else:
                f_content_base64 = f_content_raw

            if not f_content_base64:
                logger.warning(f"[{request_id}] Datei {idx+1} hat keinen Inhalt, überspringe")
                continue

            # Base64 dekodieren
            f_content_clean = f_content_base64.strip().replace('\n', '').replace('\r', '').replace(' ', '').replace('\t', '')
            valid_base64_chars = re.sub(r'[^A-Za-z0-9+/=]', '', f_content_clean)
            f_content = base64.b64decode(valid_base64_chars)
            original_file_size = len(f_content)  # Originalgröße für Usage Tracking
            logger.info(f"[{request_id}] Datei {idx+1}: {f_name} ({len(f_content)} Bytes)")
            
            # Dateityp-Konvertierung (DOCX/XLSX → PDF)
            f_extension = os.path.splitext(f_name)[1].lower()
            if f_extension in ['.docx', '.xlsx']:
                logger.info(f"[{request_id}] Konvertiere {f_name} ({f_extension}) zu PDF...")
                f_content = convert_to_pdf(f_content, f_extension, f_name, request_id)
                pdf_name = os.path.splitext(f_name)[0] + '.pdf'
                # Lokal: Konvertierte PDF in testdocs speichern
                import platform
                if platform.system() == 'Windows':
                    testdocs_dir = os.path.join(os.path.dirname(__file__), 'testdocs')
                    if os.path.isdir(testdocs_dir):
                        save_path = os.path.join(testdocs_dir, os.path.basename(pdf_name))
                        with open(save_path, 'wb') as pf:
                            pf.write(f_content)
                        logger.info(f"[{request_id}] Konvertierte PDF gespeichert: {save_path}")
                f_name = pdf_name
                f_extension = '.pdf'
            
            # PDF-Formularfelder extrahieren (werden von Azure Files API nicht gelesen)
            if f_extension == '.pdf':
                try:
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(BytesIO(f_content))
                    fields = pdf_reader.get_fields()
                    if fields:
                        form_data = []
                        for field_name, field_obj in fields.items():
                            value = field_obj.get('/V', '')
                            if value and isinstance(value, str) and value.strip():
                                form_data.append(f"{field_name}: {value}")
                        if form_data:
                            form_text = f"\n\nFORMULARFELDER aus {f_name}:\n" + "\n".join(form_data)
                            if not user_text:
                                user_text = form_text
                            else:
                                user_text = user_text + form_text
                            logger.info(f"[{request_id}] {len(form_data)} Formularfelder aus {f_name} extrahiert")
                except Exception as form_err:
                    logger.warning(f"[{request_id}] Formularfeld-Extraktion fehlgeschlagen: {form_err}")

            # Temporäre Datei erstellen
            with tempfile.NamedTemporaryFile(delete=False, suffix=f_extension) as temp_file:
                temp_file.write(f_content)
                temp_file_path = temp_file.name
            temp_file_paths.append(temp_file_path)
            processed_file_names.append(f_name)
            processed_file_sizes.append(original_file_size)  # Originalgröße speichern
            logger.info(f"[{request_id}] Temporäre Datei: {temp_file_path}")
        
        # Azure AI Client (gecacht - kein erneutes Token-Holen/Client-Erstellen)
        import time as _time
        t_start = _time.time()

        _, openai_client = get_clients(endpoint)
        agent = get_agent(endpoint, agent_name)
        t_client = _time.time()
        logger.info(f"[{request_id}] Client+Agent in {t_client - t_start:.1f}s (gecacht)")

        # Dateien hochladen
        if temp_file_paths:
            logger.info(f"[{request_id}] Lade {len(temp_file_paths)} Datei(en) zu Azure hoch...")
            for temp_path in temp_file_paths:
                with open(temp_path, "rb") as f:
                    file = openai_client.files.create(file=f, purpose="assistants")
                    uploaded_files.append(file)
                    logger.info(f"[{request_id}] Datei hochgeladen: {file.id}")
            t_upload = _time.time()
            logger.info(f"[{request_id}] Upload in {t_upload - t_client:.1f}s")
        
        # Agent-Anfrage
        logger.info(f"[{request_id}] Sende Anfrage an Agent...")
        agent_text = user_text if user_text else "Bitte analysiere die anliegende Datei."

        # Wissensbasis (Vector Store) via Search einbinden (Tools sind mit Agent-Reference nicht erlaubt)
        knowledge_context = ""
        if vector_store_ids:
            vector_query = user_text if user_text else "Bitte finde relevante Informationen zur Anfrage."
            snippets = []
            for vs_id in vector_store_ids:
                try:
                    search_results = openai_client.vector_stores.search(
                        vector_store_id=vs_id,
                        query=vector_query,
                        max_num_results=5,
                    )
                    for result in search_results:
                        content_texts = [c.text for c in getattr(result, "content", []) if getattr(c, "type", "") == "text"]
                        combined_text = " ".join(content_texts).strip()
                        if combined_text:
                            if len(combined_text) > 1200:
                                combined_text = combined_text[:1200] + "..."
                            snippets.append(f"[{result.filename} | score={result.score:.3f}] {combined_text}")
                except Exception as search_error:
                    logger.warning(f"[{request_id}] Vector Store Search fehlgeschlagen ({vs_id}): {search_error}")
            if snippets:
                knowledge_context = "WISSENSBASIS-AUSZÜGE:\n" + "\n\n".join(snippets[:10])
                logger.info(f"[{request_id}] Vector Store(s) genutzt: {vector_store_ids}")

        if uploaded_files:
            content = []
            for file in uploaded_files:
                content.append({"type": "input_file", "file_id": file.id})
            if knowledge_context:
                content.append({"type": "input_text", "text": knowledge_context})
            content.append({"type": "input_text", "text": agent_text})
        else:
            content = []
            if knowledge_context:
                content.append({"type": "input_text", "text": knowledge_context})
            content.append({"type": "input_text", "text": agent_text})

        # Retry-Logic für Rate Limits und Connection-Errors
        t_agent_start = _time.time()
        max_retries = 3
        for attempt in range(max_retries):
            try:
                request_kwargs = {
                    "input": [{"role": "user", "content": content}],
                    "extra_body": {"agent": {"name": agent.name, "type": "agent_reference"}},
                }
                response = openai_client.responses.create(**request_kwargs)
                t_agent_end = _time.time()
                logger.info(f"[{request_id}] Agent-Antwort in {t_agent_end - t_agent_start:.1f}s")
                logger.info(f"[{request_id}] Gesamtdauer bisher: {t_agent_end - t_start:.1f}s")
                break
            except Exception as api_error:
                error_str = str(api_error)
                if "429" in error_str or "Too Many Requests" in error_str:
                    if attempt < max_retries - 1:
                        wait_time = 5 * (attempt + 1)
                        logger.warning(f"[{request_id}] Rate Limit, warte {wait_time}s...")
                        sleep(wait_time)
                    else:
                        raise Exception("Azure OpenAI Rate Limit erreicht") from api_error
                elif "Connection" in error_str or "disconnected" in error_str or "RemoteProtocolError" in error_str:
                    if attempt < max_retries - 1:
                        logger.warning(f"[{request_id}] Connection-Error (Versuch {attempt+1}/{max_retries}), erstelle neuen Client...")
                        invalidate_clients(endpoint)
                        _, openai_client = get_clients(endpoint)
                        agent = get_agent(endpoint, agent_name)
                        # Dateien neu hochladen (alte file_ids sind an den alten Client gebunden)
                        if uploaded_files:
                            logger.info(f"[{request_id}] Lade Dateien erneut hoch...")
                            uploaded_files.clear()
                            for temp_path in temp_file_paths:
                                with open(temp_path, "rb") as f:
                                    file = openai_client.files.create(file=f, purpose="assistants")
                                    uploaded_files.append(file)
                            # Content mit neuen file_ids aktualisieren
                            content = []
                            for file in uploaded_files:
                                content.append({"type": "input_file", "file_id": file.id})
                            if knowledge_context:
                                content.append({"type": "input_text", "text": knowledge_context})
                            content.append({"type": "input_text", "text": agent_text})
                        sleep(2)
                    else:
                        raise
                else:
                    raise
        
        # Token-Nutzung extrahieren
        try:
            if hasattr(response, 'usage') and response.usage:
                tokens_input = getattr(response.usage, 'input_tokens', 0) or getattr(response.usage, 'prompt_tokens', 0) or 0
                tokens_output = getattr(response.usage, 'output_tokens', 0) or getattr(response.usage, 'completion_tokens', 0) or 0
                logger.info(f"[{request_id}] Tokens: Input={tokens_input}, Output={tokens_output}")
            else:
                logger.info(f"[{request_id}] Token-Nutzung nicht verfügbar in Response")
        except Exception as token_error:
            logger.warning(f"[{request_id}] Fehler beim Extrahieren der Token-Nutzung: {token_error}")
        
        # Response parsen (Code-Fences entfernen, falls vorhanden)
        try:
            output_text = response.output_text.strip() if response.output_text else ""
            if output_text.startswith("```"):
                output_text = output_text.strip("`")
                if output_text.lstrip().lower().startswith("json"):
                    output_text = output_text.lstrip()[4:].strip()
            response_data = json.loads(output_text)
        except json.JSONDecodeError:
            response_data = response.output_text
        
        # Cleanup
        for uploaded_file in uploaded_files:
            try:
                openai_client.files.delete(uploaded_file.id)
            except:
                pass
        
        # Usage Tracking - Erfolg
        log_usage(
            customer_name=customer_name,
            request_id=request_id,
            file_names=processed_file_names,
            file_sizes=processed_file_sizes,
            success=True,
            tokens_input=tokens_input,
            tokens_output=tokens_output,
            action=action
        )
        
        # Ergebnis
        result = {
            "success": True,
            "files_processed": len(uploaded_files),
            "file_names": processed_file_names,
            "response": response_data,
            "usage": {
                "tokens_input": tokens_input,
                "tokens_output": tokens_output,
                "tokens_total": tokens_input + tokens_output
            }
        }
        
        logger.info(f"[{request_id}] Analyse erfolgreich abgeschlossen")
        return result
    
    except Exception as e:
        # Usage Tracking - Fehler
        log_usage(
            customer_name=customer_name,
            request_id=request_id,
            file_names=processed_file_names,
            file_sizes=processed_file_sizes,
            success=False,
            tokens_input=tokens_input,
            tokens_output=tokens_output,
            action=action,
            error_message=str(e)
        )
        raise  # Exception weiterwerfen
        
    finally:
        # Temporäre Dateien löschen
        for temp_path in temp_file_paths:
            if os.path.exists(temp_path):
                os.unlink(temp_path)

@app.route('/analyze', methods=['POST'])
def analyze_file():
    """
    Endpoint für Power Automate
    Erwartet JSON mit:
    {
        "api_key": "customer-key",  # Erforderlich für Authentifizierung
        "action": "analyze-document",  # Erforderlich: Welche Aktion ausgeführt werden soll
        "user_text": "Optional: Zusätzlicher Text (plain-text oder HTML)",
        
        # Datei-Optionen (optional, 0-N Dateien):
        "files": [
            {"file_content": "base64-...", "file_name": "datei1.pdf"},
            {"file_content": "base64-...", "file_name": "datei2.pdf"}
        ],
        "vector_store_id": "vs_...",  # Optional: einzelner Vector Store
        "vector_store_ids": ["vs_...", "vs_..."]  # Optional: mehrere Vector Stores
        
        # Rückwärtskompatibilität (einzelne Datei):
        # "file_content": "base64-encoded-file-content",
        # "file_name": "dateiname.pdf"
    }
    
    Optional (überschreibt automatisches Mapping):
    {
        "endpoint": "https://...",  # Überschreibt automatische Bestimmung
        "agent": "agent-name"  # Überschreibt automatische Bestimmung
    }
    """
    request_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    logger.info(f"\n{'='*60}")
    logger.info(f"[{request_id}] Neue SYNC Anfrage empfangen")
    logger.info(f"{'='*60}")
    
    try:
        if not request.is_json:
            return jsonify({"error": "Request muss JSON sein"}), 400
        
        data = request.json
        logger.info(f"[{request_id}] JSON Keys: {list(data.keys())}")
        
        # Rufe interne Funktion auf
        result = analyze_document_internal(data, request_id)
        
        logger.info(f"[{request_id}] [OK] Anfrage erfolgreich abgeschlossen")
        logger.info(f"{'='*60}\n")
        return jsonify(result), 200
        
    except ValueError as ve:
        # Validierungsfehler (400)
        logger.error(f"[{request_id}] Validierungsfehler: {ve}")
        logger.info(f"{'='*60}\n")
        return jsonify({"success": False, "error": str(ve)}), 400
        
    except Exception as e:
        # Serverfehler (500)
        logger.error(f"[{request_id}] [ERROR] Fehler: {str(e)}", exc_info=True)
        logger.info(f"{'='*60}\n")
        
        error_message = str(e)
        http_status = 500
        
        if "Failed to invoke the Azure CLI" in error_message or "CredentialUnavailableError" in error_message:
            error_message = "Azure Authentifizierung fehlgeschlagen. Bitte führen Sie 'az login' aus."
            http_status = 401
        elif "429" in error_message or "Too Many Requests" in error_message or "Rate Limit" in error_message:
            error_message = "Azure OpenAI Rate Limit erreicht. Zu viele Anfragen in kurzer Zeit. Bitte warten Sie 1-2 Minuten."
            http_status = 429
        elif "Autorisierung fehlgeschlagen" in error_message:
            http_status = 403
        
        return jsonify({"success": False, "error": error_message}), http_status

@app.route('/analyze-async', methods=['POST'])
def analyze_file_async():
    """
    Asynchroner Endpoint für Power Automate (für lange Verarbeitungszeiten)
    
    Startet die Analyse im Hintergrund und gibt sofort eine task_id zurück.
    
    Request Body: Gleich wie /analyze
    
    Response:
    {
        "task_id": "uuid",
        "status": "pending",
        "status_url": "/status/<task_id>"
    }
    """
    request_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    logger.info(f"\n{'='*60}")
    logger.info(f"[{request_id}] Neue ASYNC Anfrage empfangen")
    logger.info(f"{'='*60}")
    
    try:
        if not request.is_json:
            return jsonify({"error": "Request muss JSON sein"}), 400
        
        data = request.json
        
        # Generiere eindeutige Task-ID
        task_id = str(uuid.uuid4())
        
        # Erstelle Job
        create_job(task_id)
        logger.info(f"[{request_id}] Task erstellt: {task_id}")
        
        # Windows: Thread (spawn-basiertes multiprocessing verursacht Connection-Probleme)
        # Linux/Azure: Eigener Prozess (eigene GIL, kein Konkurrenzkampf mit Request-Handling)
        import platform
        if platform.system() == 'Windows':
            thread = threading.Thread(
                target=process_async_job,
                args=(task_id, data),
                daemon=True
            )
            thread.start()
            logger.info(f"[{task_id}] Background-Thread gestartet (Windows)")
        else:
            process = multiprocessing.Process(
                target=process_async_job,
                args=(task_id, data),
                daemon=True
            )
            process.start()
            logger.info(f"[{task_id}] Background-Prozess gestartet (Linux)")
        
        # Cleanup alte Jobs
        cleanup_old_jobs()
        
        return jsonify({
            "task_id": task_id,
            "status": "pending",
            "status_url": f"/status/{task_id}",
            "message": "Analyse wurde gestartet. Verwenden Sie den status_url um den Fortschritt zu prüfen."
        }), 202
        
    except Exception as e:
        logger.error(f"[{request_id}] Fehler beim Starten des Async-Jobs: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/status/<task_id>', methods=['GET'])
def get_status(task_id):
    """
    Statusabfrage für asynchrone Jobs
    
    Response bei pending/processing:
    {
        "task_id": "uuid",
        "status": "processing",
        "progress": 50,
        "created_at": "...",
        "started_at": "..."
    }
    
    Response bei completed:
    {
        "task_id": "uuid",
        "status": "completed",
        "progress": 100,
        "created_at": "...",
        "completed_at": "...",
        "result": { ... }  # Gleiche Struktur wie synchroner /analyze Endpoint
    }
    
    Response bei failed:
    {
        "task_id": "uuid",
        "status": "failed",
        "error": "Fehlermeldung"
    }
    """
    logger.info(f"Status-Abfrage für Task: {task_id}")
    
    job = get_job(task_id)
    
    if not job:
        return jsonify({"error": "Task nicht gefunden"}), 404
    
    # Basis-Response
    response = {
        "task_id": job['task_id'],
        "status": job['status'],
        "progress": job['progress'],
        "created_at": job['created_at']
    }
    
    # Erweitere Response basierend auf Status
    if job['status'] in ['processing', 'completed']:
        if job['started_at']:
            response['started_at'] = job['started_at']
    
    if job['status'] == 'completed':
        response['completed_at'] = job['completed_at']
        response['result'] = job['result']
    
    if job['status'] == 'failed':
        response['completed_at'] = job['completed_at']
        response['error'] = job['error']
    
    return jsonify(response), 200

@app.route('/test-conversion', methods=['POST'])
def test_conversion():
    """Test endpoint für DOCX/XLSX Konvertierung"""
    try:
        data = request.get_json()
        file_content = data.get('file_content')
        file_name = data.get('file_name', 'test.docx')
        
        if not file_content:
            return jsonify({"error": "file_content fehlt"}), 400
        
        # Decode file
        file_bytes = base64.b64decode(file_content)
        
        # Detect file type
        file_ext = os.path.splitext(file_name)[1].lower()
        
        result = {
            "original_file": file_name,
            "file_size": len(file_bytes),
            "file_type": file_ext
        }
        
        # Try conversion
        if file_ext in ['.docx', '.xlsx']:
            try:
                logger.info(f"Test-Konvertierung: {file_name} ({len(file_bytes)} bytes)")
                pdf_bytes = convert_to_pdf(file_bytes, file_ext)
                
                if pdf_bytes:
                    result["conversion"] = "success"
                    result["pdf_size"] = len(pdf_bytes)
                    
                    # Try to extract text with PyPDF2
                    try:
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                        result["extracted_text_length"] = len(text.strip())
                        result["extracted_text_preview"] = text.strip()[:200]
                    except Exception as e:
                        result["text_extraction_error"] = str(e)
                else:
                    result["conversion"] = "failed - returned None"
            except Exception as e:
                result["conversion"] = "failed"
                result["error"] = str(e)
                result["traceback"] = traceback.format_exc()
        else:
            result["conversion"] = "skipped - not docx/xlsx"
        
        return jsonify(result), 200
        
    except Exception as e:
        logger.error(f"Test-Conversion Error: {e}\n{traceback.format_exc()}")
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/health', methods=['GET'])
def health():
    """Health Check Endpoint"""
    logger.info("Health Check aufgerufen")
    return jsonify({"status": "ok"}), 200

@app.route('/usage', methods=['GET'])
def get_usage():
    """
    Abrechnungs-Endpoint: Gibt die Nutzungsstatistiken zurück
    
    Query-Parameter (alle optional):
    - customer: Filtert nach Kundenname
    - start_date: Startdatum (YYYY-MM-DD)
    - end_date: Enddatum (YYYY-MM-DD)
    - format: "summary" (nur Zusammenfassung) oder "full" (mit allen Einträgen)
    
    Beispiel: /usage?customer=Fameline&start_date=2026-01-01&end_date=2026-01-31&format=summary
    """
    customer = request.args.get('customer')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    output_format = request.args.get('format', 'full')
    
    logger.info(f"Usage-Abfrage: customer={customer}, start={start_date}, end={end_date}, format={output_format}")
    
    result = get_usage_summary(customer_name=customer, start_date=start_date, end_date=end_date)
    
    if output_format == 'summary':
        # Nur Zusammenfassung ohne Einzeleinträge
        return jsonify({"summary": result.get("summary", {})}), 200
    else:
        return jsonify(result), 200

@app.route('/usage/export', methods=['GET'])
def export_usage():
    """
    Exportiert die Nutzungsdaten als CSV
    
    Query-Parameter (alle optional):
    - customer: Filtert nach Kundenname
    - start_date: Startdatum (YYYY-MM-DD)
    - end_date: Enddatum (YYYY-MM-DD)
    """
    from flask import Response
    
    customer = request.args.get('customer')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    
    result = get_usage_summary(customer_name=customer, start_date=start_date, end_date=end_date)
    entries = result.get("entries", [])
    
    # CSV erstellen
    csv_lines = ["Datum,Uhrzeit,Kunde,Aktion,Erfolg,Dateien,Dateigröße (Bytes),Tokens Input,Tokens Output,Tokens Gesamt,Fehler"]
    
    for entry in entries:
        files_str = "; ".join([f.get("name", "") for f in entry.get("files", [])])
        tokens = entry.get("tokens", {})
        csv_lines.append(
            f'{entry.get("date", "")},{entry.get("time", "")},"{entry.get("customer", "")}",'
            f'"{entry.get("action", "")}",{entry.get("success", False)},"{files_str}",'
            f'{entry.get("total_file_size_bytes", 0)},{tokens.get("input", 0)},'
            f'{tokens.get("output", 0)},{tokens.get("total", 0)},"{entry.get("error", "") or ""}"'
        )
    
    csv_content = "\n".join(csv_lines)
    
    # Dateiname generieren
    filename = f"usage_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    if customer:
        filename = f"usage_{customer}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    
    return Response(
        csv_content,
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={filename}'}
    )

# ============================================
# Admin Endpoints für Kundenverwaltung
# ============================================

@app.route('/admin', methods=['GET'])
@require_auth  # Nur Admin-UI ist geschützt
def admin_ui():
    """Serve Admin UI"""
    admin_ui_path = os.path.join(os.path.dirname(__file__), 'ui', 'admin_ui.html')
    if not os.path.exists(admin_ui_path):
        logger.error(f"Admin UI nicht gefunden: {admin_ui_path}")
        return jsonify({
            "error": "Admin UI nicht gefunden",
            "path": admin_ui_path
        }), 404

    try:
        with open(admin_ui_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Fehler beim Laden der Admin UI: {e}")
        return jsonify({
            "error": "Admin UI konnte nicht geladen werden",
            "details": str(e)
        }), 500

@app.route('/admin/customers', methods=['GET'])
@require_auth  # Nur Admin-Endpoints sind geschützt
def get_customers():
    """Liste aller Kunden zurückgeben"""
    try:
        with open(CUSTOMER_CONFIG_FILE, 'r', encoding='utf-8') as f:
            config = json.load(f)
        return jsonify(config), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/admin/customers', methods=['POST'])
@require_auth  # Nur Admin-Endpoints sind geschützt
def create_customer():
    """Neuen Kunden anlegen"""
    try:
        data = request.json
        api_key = data.get('api_key')
        customer = data.get('customer')
        
        if not api_key or not customer:
            return jsonify({"error": "api_key und customer sind erforderlich"}), 400
        
        # Config laden
        with open(CUSTOMER_CONFIG_FILE, 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        # Prüfen ob Key schon existiert
        if api_key in config.get('customers', {}):
            return jsonify({"error": "API-Key existiert bereits"}), 400
        
        # Kunden hinzufügen
        if 'customers' not in config:
            config['customers'] = {}
        config['customers'][api_key] = customer
        
        # Speichern
        with open(CUSTOMER_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2, ensure_ascii=False)
        
        # Config neu laden
        load_customer_config()
        
        return jsonify({"success": True, "message": "Kunde erfolgreich angelegt"}), 201
        
    except Exception as e:
        logger.error(f"Error creating customer: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/admin/customers', methods=['PUT'])
@require_auth  # Nur Admin-Endpoints sind geschützt
def update_customer():
    """Bestehenden Kunden aktualisieren"""
    try:
        data = request.json
        api_key = data.get('api_key')
        customer = data.get('customer')
        
        if not api_key or not customer:
            return jsonify({"error": "api_key und customer sind erforderlich"}), 400
        
        # Config laden
        with open(CUSTOMER_CONFIG_FILE, 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        # Prüfen ob Key existiert
        if api_key not in config.get('customers', {}):
            return jsonify({"error": "API-Key nicht gefunden"}), 404
        
        # Kunden aktualisieren
        config['customers'][api_key] = customer
        
        # Speichern
        with open(CUSTOMER_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2, ensure_ascii=False)
        
        # Config neu laden
        load_customer_config()
        
        return jsonify({"success": True, "message": "Kunde erfolgreich aktualisiert"}), 200
        
    except Exception as e:
        logger.error(f"Error updating customer: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/admin/customers', methods=['DELETE'])
@require_auth  # Nur Admin-Endpoints sind geschützt
def delete_customer():
    """Kunden löschen"""
    try:
        data = request.json
        api_key = data.get('api_key')
        
        if not api_key:
            return jsonify({"error": "api_key ist erforderlich"}), 400
        
        # Config laden
        with open(CUSTOMER_CONFIG_FILE, 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        # Prüfen ob Key existiert
        if api_key not in config.get('customers', {}):
            return jsonify({"error": "API-Key nicht gefunden"}), 404
        
        # Kunden löschen
        del config['customers'][api_key]
        
        # Speichern
        with open(CUSTOMER_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2, ensure_ascii=False)
        
        # Config neu laden
        load_customer_config()
        
        return jsonify({"success": True, "message": "Kunde erfolgreich gelöscht"}), 200
        
    except Exception as e:
        logger.error(f"Error deleting customer: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    logger.info("=" * 60)
    logger.info("Server startet auf http://0.0.0.0:5000")
    logger.info("Endpoints:")
    logger.info("  - POST /analyze       - Datei-Analyse (synchron, max 900s)")
    logger.info("  - POST /analyze-async - Datei-Analyse (asynchron, für Power Automate)")
    logger.info("  - GET  /status/<id>   - Status-Abfrage für async Jobs")
    logger.info("  - GET  /usage         - Nutzungsstatistiken (Abrechnung)")
    logger.info("  - GET  /usage/export  - CSV-Export der Nutzungsdaten")
    logger.info("  - GET  /admin         - Admin-UI für Kundenverwaltung")
    logger.info("  - GET  /health        - Health Check")
    logger.info("=" * 60)
    # Für Produktion sollte ein WSGI-Server wie gunicorn verwendet werden
    app.run(host='0.0.0.0', port=5000, debug=False)
